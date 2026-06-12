#!/usr/bin/env python3
"""构造 fixtures 语料(P 流)。

fixtures 不入库(可重建);本脚本据 ``fixtures_sources.csv`` 真下载外规公开法规。

子命令(flags):
- ``--download``    : 下载外规 PDF(证监会/交易所公开法规),仅收**文本层**(pypdf 校验,防图片版)
- ``--gen-internal``: 生成自拟内规 docx(P2,待实现)

合规:政府/交易所公开信息,可直接用、无需脱敏(SPEC 决策6)。
"""

from __future__ import annotations

import argparse
import csv
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
BATCH01 = ROOT / "fixtures" / "batch01"
SOURCES = Path(__file__).resolve().parent / "fixtures_sources.csv"
_UA = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
)
MIN_TEXT_CHARS = 200  # 文本层判据:前 5 页抽出的非空文本字符数下限


def _curl(url: str, dest: Path) -> None:
    subprocess.run(
        ["curl", "-sSL", "-A", _UA, "-m", "90", "--retry", "3", "-o", str(dest), url], check=True
    )


def _text_layer_pages(pdf_path: Path) -> tuple[bool, int]:
    reader = PdfReader(str(pdf_path))
    text = "".join((p.extract_text() or "") for p in reader.pages[:5])
    return len(text.strip()) >= MIN_TEXT_CHARS, len(reader.pages)


def download() -> int:
    BATCH01.mkdir(parents=True, exist_ok=True)
    with SOURCES.open(encoding="utf-8-sig") as f:
        rows = list(csv.DictReader(f))
    ok = 0
    for r in rows:
        dest = BATCH01 / r["filename"]
        print(f"↓ {r['title']} → {dest.name}")
        try:
            _curl(r["url"], dest)
        except subprocess.CalledProcessError as e:
            print(f"  ✗ 下载失败: {e}")
            continue
        if not dest.exists() or dest.stat().st_size == 0:
            print("  ✗ 空文件,丢弃")
            dest.unlink(missing_ok=True)
            continue
        text_ok, npages = _text_layer_pages(dest)
        if not text_ok:
            print(f"  ✗ 非文本层(疑似图片版),丢弃: {dest.name}")
            dest.unlink(missing_ok=True)
            continue
        print(f"  ✓ {npages} 页,文本层 OK")
        ok += 1
    print(f"\n外规下载完成:{ok}/{len(rows)} 篇 → {BATCH01}")
    return ok


# ── P2:自拟内规 docx(条款号写字面文本,规避 R5b 自动编号)──────────────────
_CN_SEQ = [
    "一", "二", "三", "四", "五", "六", "七", "八",
    "九", "十", "十一", "十二", "十三", "十四", "十五", "十六",
]

_REPORT_ITEMS = [
    "涉及金额超过五十万元的对外投资、资产处置、产权交易或者大额资金往来等重要经济事项",
    "可能对本单位声誉、形象或者财务状况产生重大不利影响的诉讼、仲裁、行政处罚等事项",
    "重要业务信息系统发生较长时间中断、重要数据丢失或者泄露、重大网络安全等事件",
    "因自然灾害、事故灾难、公共卫生等突发事件造成人员伤亡或者较大财产损失的情形",
    "本单位领导班子成员因公因私出国境、离岗学习以及连续较长时间休假的有关情况",
    "涉及群体性事件、重要信访突出问题或者可能引发较大负面舆情的敏感事项",
    "重大合同的订立、重大条款变更、提前解除以及合同履行过程中发生的重大争议",
    "重要固定资产或者无形资产的抵押、质押、对外担保以及向外部单位提供大额借款",
    "内部机构设置调整、人员编制变化以及重要管理岗位负责人的选拔任用与变动",
    "年度预算的重大调整、较大金额预算外支出以及重大投资、融资计划的安排",
    "审计、巡视、专项检查中发现的重大问题以及问题整改落实的进展和结果情况",
    "重大改革举措的研究、制定与组织实施过程中的重要情况和阶段性进展",
    "与监管部门、上级单位之间就重大事项进行的重要沟通、协调与请示报告",
    "涉及国有资产产权登记、评估、转让以及收益管理中的重大问题和异常情况",
    "上级机关交办、督办的重要事项以及落实重大决策部署中的关键进展情况",
    "其他依照有关法律法规、上级规定和本单位制度应当及时报告的重大事项",
]


def _build(title: str, paragraphs: list[str]) -> DocxDocument:
    d = Document()
    d.add_paragraph(title)
    for p in paragraphs:
        d.add_paragraph(p)
    return d


def _doc_baoxiao() -> DocxDocument:
    return _build(
        "XX单位费用报销管理办法",
        [
            "第一章 总则",
            "第一条 为规范本单位费用报销管理,加强财务监督,根据国家有关财务制度,制定本办法。",
            "第二条 本办法适用于本单位各部门及全体工作人员的费用报销活动。",
            "第二章 报销范围与标准",
            "第三条 可报销费用包括差旅费、办公费、会议费、培训费等与公务直接相关的支出。",
            "第四条 各项费用报销应当符合本单位规定的开支标准,超标准部分原则上不予报销。",
            "第四条之一 因特殊事由确需超标准开支的,应当事前书面报分管领导审批。",
            "第三章 报销流程",
            "第五条 报销人应如实填写报销单,附合规发票及凭证,经部门负责人审核后报财务部门。",
            "第六条 财务部门应当在收到完整报销申请后五个工作日内完成审核与支付。",
            "第四章 附则",
            "第七条 本办法自发布之日起施行,由财务部门负责解释。",
        ],
    )


def _doc_hetong() -> DocxDocument:
    return _build(
        "XX单位合同审批管理办法",
        [
            "第一章 总则",
            "第一条 为规范合同管理,防范法律风险,根据有关法律法规,制定本办法。",
            "第二条 本办法所称合同,是指本单位对外订立的各类协议。",
            "第二章 审批权限",
            "第一节 一般合同",
            "第三条 金额在五十万元以下的一般合同,由分管副职审批。",
            "第二节 重大合同",
            "第四条 金额在五十万元以上或者涉及重大权利义务的合同,由主要负责人审批。",
            "第三章 附则",
            "第五条 本办法自发布之日起施行。",
        ],
    )


def _doc_yinzhang() -> DocxDocument:
    return _build(
        "XX单位印章使用管理办法",
        [
            "第一章 总则",
            "第一条 为规范印章使用,防止印章管理风险,制定本办法。",
            "第二条 本办法适用于本单位公章、合同专用章、财务专用章的管理与使用。",
            "第二章 用印管理",
            "第三条 用印须填写用印申请单,经审批后方可加盖。",
            "第四条 印章由专人保管,保管人不得擅自用印。",
            "第五条 重要文件用印应当登记备查。",
            "第三章 附则",
            "第六条 本办法自发布之日起施行。",
        ],
    )


def _doc_quanxian() -> DocxDocument:
    d = _build(
        "XX单位审批权限管理办法",
        [
            "第一条 为明确各级审批权限,规范审批行为,制定本办法。",
            "第二条 各类事项的审批权限,按照下表执行:",
        ],
    )
    rows = [
        ("事项", "金额区间", "审批人", "备注"),
        ("办公用品采购", "1万元以下", "部门负责人", "据实报销"),
        ("办公用品采购", "1万元至5万元", "分管副职", "需比价"),
        ("固定资产购置", "5万元至20万元", "主要负责人", "需论证"),
        ("固定资产购置", "20万元以上", "领导班子集体决策", "需招标"),
        ("对外捐赠", "10万元以下", "分管副职", ""),
        ("对外捐赠", "10万元以上", "主要负责人", ""),
        ("大额资金支付", "50万元以上", "领导班子集体决策", "需专项报告"),
    ]
    table = d.add_table(rows=len(rows), cols=4)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.cell(r, c).text = val
    d.add_paragraph("第三条 本办法自发布之日起施行。")
    return d


def _doc_baogao() -> DocxDocument:
    items = "；".join(f"（{cn}）{it}" for cn, it in zip(_CN_SEQ, _REPORT_ITEMS, strict=False))
    prefix = "第二条 本单位发生下列重大事项的,有关部门应当在二十四小时内书面报告:"
    long_clause = prefix + items + "。"
    return _build(
        "XX单位重大事项报告办法",
        [
            "第一条 为规范重大事项报告,加强内部管理,制定本办法。",
            long_clause,  # 超长条款:单段 >600 字
            "第三条 本办法自发布之日起施行。",
        ],
    )


def _doc_tongzhi() -> DocxDocument:
    return _build(
        "关于规范办公用品采购的通知",
        [
            "各部门:",
            "为规范办公用品采购,厉行节约,现就有关事项通知如下。",
            "第一条 办公用品采购应当编制年度计划,统一归口管理。",
            "第二条 单次采购金额在一万元以下的,由部门负责人审批。",
            "第三条 本通知自发布之日起执行。",
            "XX单位办公室",
        ],
    )


#: (元数据, builder)。元数据供 P4 生成 manifest;均 P-INT、issuer=INTERNAL、密级=内部。
INTERNAL_DOCS = [
    ({"filename": "int_baoxiao.docx", "title": "XX单位费用报销管理办法",
      "doc_number": "XX发〔2024〕1号", "biz_domain": "EXPENSE"}, _doc_baoxiao),
    ({"filename": "int_hetong.docx", "title": "XX单位合同审批管理办法",
      "doc_number": "XX发〔2024〕2号", "biz_domain": "CONTRACT"}, _doc_hetong),
    ({"filename": "int_yinzhang.docx", "title": "XX单位印章使用管理办法",
      "doc_number": "XX发〔2024〕3号", "biz_domain": "APPROVAL"}, _doc_yinzhang),
    ({"filename": "int_quanxian.docx", "title": "XX单位审批权限管理办法",
      "doc_number": "XX发〔2024〕4号", "biz_domain": "APPROVAL"}, _doc_quanxian),
    ({"filename": "int_baogao.docx", "title": "XX单位重大事项报告办法",
      "doc_number": "XX发〔2024〕5号", "biz_domain": "FINANCE"}, _doc_baogao),
    ({"filename": "int_tongzhi.docx", "title": "关于规范办公用品采购的通知",
      "doc_number": "XX办〔2024〕6号", "biz_domain": "FINANCE"}, _doc_tongzhi),
]


def gen_internal() -> int:
    BATCH01.mkdir(parents=True, exist_ok=True)
    for meta, builder in INTERNAL_DOCS:
        builder().save(str(BATCH01 / meta["filename"]))
        print(f"  ✓ {meta['title']} → {meta['filename']}")
    _verify_internal()
    print(f"\n内规生成完成:{len(INTERNAL_DOCS)} 件 → {BATCH01}")
    return len(INTERNAL_DOCS)


def _verify_internal() -> None:
    paths = sorted(BATCH01.glob("int_*.docx"))
    assert len(paths) == 6, f"内规应 6 件,实际 {len(paths)}"
    has_zhiyi = has_big_table = has_long = has_no_chapter = False
    chapter_docs = 0
    for p in paths:
        d = Document(str(p))
        texts = [par.text for par in d.paragraphs]
        # 字面条款号:存在以"第X条"开头的段落(非自动编号)
        assert any(re.match(r"^第.{1,8}条", t.strip()) for t in texts), f"{p.name} 无字面条款号"
        if any(re.search(r"第.{1,6}条之一", t) for t in texts):
            has_zhiyi = True
        if any(len(t) > 600 for t in texts):
            has_long = True
        if any(re.match(r"^第.{1,6}章", t.strip()) for t in texts):
            chapter_docs += 1
        else:
            has_no_chapter = True
        if any(len(tbl.rows) >= 6 for tbl in d.tables):
            has_big_table = True
    assert has_zhiyi, "缺『第X条之一』插入条"
    assert has_big_table, "缺大表格(≥6 行)"
    assert has_long, "缺超长条款(单段 >600 字)"
    assert has_no_chapter, "缺无章直条件(短通知)"
    assert chapter_docs >= 3, f"标准章节条件应 ≥3,实际 {chapter_docs}"


def main() -> None:
    ap = argparse.ArgumentParser(description="构造 fixtures 语料")
    ap.add_argument("--download", action="store_true", help="下载外规 PDF(仅文本层)")
    ap.add_argument("--gen-internal", action="store_true", help="生成自拟内规 docx(P2)")
    args = ap.parse_args()

    if args.download:
        if download() < 3:
            sys.exit("✗ 外规真下载不足 3 篇,P1 验收不通过")
    if args.gen_internal:
        if gen_internal() < 6:
            sys.exit("✗ 内规生成不足 6 件,P2 验收不通过")
    if not (args.download or args.gen_internal):
        ap.print_help()


if __name__ == "__main__":
    main()
