"""s1_parse 集成测试(连真 PG + ObjectStore + soffice;不可达自动 skip)。"""

import io

import pytest
from docx import Document as Docx
from PIL import Image, ImageDraw
from sqlalchemy import delete, select, text
from ulid import ULID

from pipeline.config import load_config
from pipeline.index.object_store import ObjectStore
from pipeline.index.pg_io import PgIO
from pipeline.index.pg_models import Document, DocVersion, ImportBatch, PipelineEvent
from pipeline.parsing.rendition import render_pdf
from pipeline.stage_base import StageContext
from pipeline.stages import s1_parse as s1
from pipeline.states import PipelineState as PS


@pytest.fixture
def pg():
    io_ = PgIO.from_config(load_config())
    try:
        with io_.session() as s:
            s.execute(text("select 1"))
    except Exception:
        pytest.skip("PG 不可达(demo up 未起)")
    return io_


@pytest.fixture
def env(pg, tmp_path):
    ctx = StageContext(config=load_config(), object_store=ObjectStore(tmp_path / "obj"), db=pg)
    batches: list[str] = []
    yield ctx, batches
    with pg.session() as s:
        dvs = list(s.scalars(select(DocVersion).where(DocVersion.batch_id.in_(batches or [""]))))
        dvids = [d.doc_version_id for d in dvs]
        lids = {d.logical_id for d in dvs}
        if dvids:
            s.execute(delete(PipelineEvent).where(PipelineEvent.doc_version_id.in_(dvids)))
            s.execute(delete(DocVersion).where(DocVersion.doc_version_id.in_(dvids)))
        if lids:
            s.execute(delete(Document).where(Document.logical_id.in_(lids)))
        if batches:
            s.execute(delete(ImportBatch).where(ImportBatch.batch_id.in_(batches)))


def _make_doc(ctx, fmt, data, *, corpus="P-INT") -> tuple[str, str]:
    bid, lid, dvid = "t_" + str(ULID()), str(ULID()), str(ULID())
    ctx.db.add(ImportBatch(batch_id=bid, source_dir="x"))
    ctx.db.add(Document(logical_id=lid, corpus_type=corpus))
    raw_key = ctx.object_store.put_raw(corpus, bid, dvid, fmt, data)
    ctx.db.add(
        DocVersion(
            doc_version_id=dvid, logical_id=lid, batch_id=bid, source_format=fmt,
            source_hash="h" + dvid[:10], raw_object_key=raw_key, source_filename=f"x.{fmt}",
            pipeline_status="PARSING",
        )
    )
    return bid, dvid


def _docx_bytes() -> bytes:
    buf = io.BytesIO()
    d = Docx()
    d.add_paragraph("某单位综合管理办法")
    d.add_paragraph("第一章 总则")
    d.add_paragraph("第一条 为加强本单位综合管理,规范各项工作流程,根据有关规定制定本办法。")
    d.add_paragraph("第二条 本办法适用于本单位各部门及全体工作人员的日常管理活动。")
    d.add_paragraph("第三条 各部门应当按照职责分工,密切配合,共同做好综合管理工作。")
    d.save(buf)
    return buf.getvalue()


def _scanned_pdf_bytes() -> bytes:
    im = Image.new("RGB", (800, 1000), "white")
    ImageDraw.Draw(im).text((50, 50), "SCANNED (image only)", fill="black")
    buf = io.BytesIO()
    im.save(buf, format="PDF")
    return buf.getvalue()


def test_start_claims_registered_to_parsing():
    # 薄 stage:REGISTERED → PARSING,纯状态翻转,不读不写、无入队/错误码(无需 PG/soffice)
    res = s1.start(StageContext(config=load_config()), "anydvid")
    assert res.next_state is PS.PARSING
    assert res.queue is None
    assert res.error_code is None


def test_docx_renders_aligns_writes_ir(env, soffice):
    ctx, batches = env
    bid, dvid = _make_doc(ctx, "docx", _docx_bytes())
    batches.append(bid)
    res = s1.run(ctx, dvid)
    assert res.next_state is PS.QC_PENDING
    assert ctx.object_store.exists_rendition(dvid)  # 渲染件落库
    ir = ctx.object_store.load_ir(dvid)
    assert ir.blocks and any(b.page is not None for b in ir.blocks)  # 页码对齐回填
    dv = ctx.db.get(DocVersion, dvid)
    assert dv.ir_object_key and dv.rendition_object_key


def test_pdf_native_pages(env, tmp_path, soffice):
    ctx, batches = env
    src = tmp_path / "s.docx"
    src.write_bytes(_docx_bytes())
    pdf = render_pdf(src, tmp_path)
    bid, dvid = _make_doc(ctx, "pdf", pdf.read_bytes())
    batches.append(bid)
    res = s1.run(ctx, dvid)
    assert res.next_state is PS.QC_PENDING
    ir = ctx.object_store.load_ir(dvid)
    assert ir.blocks and all(b.page is not None for b in ir.blocks)  # pdf 原生页码
    assert ctx.db.get(DocVersion, dvid).rendition_object_key is None  # pdf 无渲染件


def test_scanned_pdf_quarantined(env):
    ctx, batches = env
    bid, dvid = _make_doc(ctx, "pdf", _scanned_pdf_bytes())
    batches.append(bid)
    res = s1.run(ctx, dvid)
    assert res.next_state is PS.QUARANTINED
    assert res.error_code == "E202-DEMO"
    assert res.queue is not None and res.queue.queue_type == "quarantine"


def test_rendition_write_once_on_reprocess(env, soffice):
    ctx, batches = env
    bid, dvid = _make_doc(ctx, "docx", _docx_bytes())
    batches.append(bid)
    s1.run(ctx, dvid)
    first = ctx.object_store.get_rendition(dvid)
    s1.run(ctx, dvid)  # reprocess:复用渲染件,不重渲
    assert ctx.object_store.get_rendition(dvid) == first  # 写一次,字节不变
