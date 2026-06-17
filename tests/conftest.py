"""共享 fixtures。

soffice 探测:**二进制找得到 ≠ 能渲染**(profile 锁 / 缺字体库 / 静默退 0 不产出 / 超时)。
本 fixture 用与生产同一个 ``render_pdf`` 真渲一份最小 docx,任何失败即 skip 所有渲染相关
测试——把"环境不可用"与"代码回归"区分开:已知良好输入都渲染不出是环境问题(skip);
某具体测试在探测通过后仍渲染失败,才是真 bug(照常 fail)。
"""

from __future__ import annotations

import io
import shutil
from pathlib import Path

import pytest
from docx import Document as Docx
from openpyxl import Workbook, load_workbook
from ulid import ULID

from pipeline.parsing.rendition import render_pdf, soffice_bin

_MANIFEST_COLS = [
    "filename", "title", "doc_number", "issuer", "perm_tag",
    "corpus_type", "biz_domain", "issue_date", "supersedes",
]


@pytest.fixture(scope="session")
def soffice(tmp_path_factory):
    """探测 soffice 真能 docx→PDF;二进制缺失或渲染失败均 skip。返回可用的 soffice 路径。"""
    try:
        bin_ = soffice_bin()
    except RuntimeError as e:
        pytest.skip(str(e))
    d = tmp_path_factory.mktemp("soffice_probe")
    src = d / "probe.docx"
    buf = io.BytesIO()
    doc = Docx()
    doc.add_paragraph("探测渲染")
    doc.save(buf)
    src.write_bytes(buf.getvalue())
    try:
        render_pdf(src, d, timeout=60)
    except Exception as e:  # 二进制在但渲染崩:环境问题,跳过渲染相关测试(非代码缺陷)
        pytest.skip(f"soffice 存在但渲染失败(环境问题,非代码): {e}")
    return bin_


@pytest.fixture
def mini_batch():
    """从 fixtures 抽**单件** → 临时批目录(原件 + 该件 manifest 行,9 列契约不破)。

    返回 ``make(tmp_path, batch, filename) -> (batch_dir, manifest_path)``。聚焦单件(尤其两件
    外规 PDF)避免整 batch01 12 件解析慢 + 免 docx soffice;D2/D3 集成测试共用。
    """

    def _make(tmp_path: Path, batch: str, filename: str) -> tuple[Path, Path]:
        src_dir = Path("fixtures") / batch
        ws = load_workbook(src_dir / "manifest.xlsx").active
        rows = list(ws.iter_rows(values_only=True))
        hdr = list(rows[0])
        row = next(r for r in rows[1:] if r[hdr.index("filename")] == filename)
        d = tmp_path / batch
        d.mkdir()
        shutil.copy(src_dir / filename, d / filename)
        out = Workbook()
        out.active.append(hdr)
        out.active.append(list(row))
        mpath = d / "manifest.xlsx"
        out.save(mpath)
        return d, mpath

    return _make


@pytest.fixture
def ingest_index():
    """真实接入→索引:register_batch → orchestrator 到 META_REVIEW → 逐件 ``cli._approve_doc``
    (含自动 finalize)→ INDEXED。返回 ``make(pg, ctx, dir, manifest) -> (bid, dvids)``。

    走真实 CLI 路径(D2/D3 集成测试共用),覆盖人工闸放行 + 自动版本切换。
    """
    from sqlalchemy import select
    from ulid import ULID

    from common.pg_models import DocVersion
    from pipeline import cli
    from pipeline.orchestrator import Orchestrator
    from pipeline.stages.s0_register import register_batch

    def _make(pg, ctx, batch_dir: Path, manifest: Path) -> tuple[str, list[str]]:
        bid = str(ULID())
        register_batch(ctx, bid, batch_dir, manifest)
        Orchestrator(pg, ctx, cli._build_stages()).run_until_idle()
        with pg.session() as s:
            dvids = [
                d.doc_version_id
                for d in s.scalars(select(DocVersion).where(DocVersion.batch_id == bid))
            ]
        for dvid in dvids:
            cli._approve_doc(pg, ctx, dvid, "test")
        return bid, dvids

    return _make


@pytest.fixture
def unique_docx():
    """生成**唯一**内规 docx(嵌 ULID 保 SHA 不与库内既有件撞)+ 1 行 manifest。

    返回 ``make(tmp_path) -> (batch_dir, manifest)``。供需「自造可 ingest 件」的集成测试(T2/T4)避开
    走查/既有数据的 SHA 去重。条款含实质文本以稳过 QC 到 META_REVIEW。
    """

    def _make(tmp_path: Path) -> tuple[Path, Path]:
        tag = str(ULID())
        d = tmp_path / ("u_" + tag[:8])
        d.mkdir()
        fn = "uniq.docx"
        doc = Docx()
        doc.add_paragraph("第一章 总则")
        doc.add_paragraph(
            f"第一条 为加强本单位合同管理规范合同签订与履行流程根据有关规定制定本办法编号{tag}。"
        )
        doc.add_paragraph("第二条 本办法适用于本单位各部门及全体人员的合同签订与履行活动。")
        doc.add_paragraph("第三条 合同应当经法务审查并由授权人签署后方可对外签订生效。")
        doc.save(d / fn)
        wb = Workbook()
        wb.active.append(_MANIFEST_COLS)
        wb.active.append([fn, "合同管理办法", f"测试第{tag[:6]}号", "INTERNAL",
                          "内部", "P-INT", "LEGAL", None, None])
        mp = d / "manifest.xlsx"
        wb.save(mp)
        return d, mp

    return _make
