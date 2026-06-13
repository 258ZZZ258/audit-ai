import io

import pytest
from docx import Document as Docx
from PIL import Image, ImageDraw

from pipeline.ir import BlockType
from pipeline.parsing.light_parser import LightParser


def _docx_bytes() -> bytes:
    buf = io.BytesIO()
    d = Docx()
    d.add_paragraph("某单位综合管理办法")
    d.add_paragraph("第一章 总则")
    d.add_paragraph("第一条 为加强本单位综合管理,规范各项工作流程,根据有关规定制定本办法。")
    d.add_paragraph("第二条 本办法适用于本单位各部门及全体工作人员的日常管理活动。")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text, t.cell(0, 1).text = "事项", "责任部门"
    t.cell(1, 0).text, t.cell(1, 1).text = "综合协调", "办公室"
    d.add_paragraph("第三条 各部门应当按照职责分工,密切配合,共同做好综合管理工作。")
    d.add_paragraph("第四条 本办法自发布之日起施行,由办公室负责解释。")
    d.save(buf)
    return buf.getvalue()


def _scanned_pdf_bytes() -> bytes:
    im = Image.new("RGB", (800, 1000), "white")
    ImageDraw.Draw(im).text((50, 50), "SCANNED (image only)", fill="black")
    buf = io.BytesIO()
    im.save(buf, format="PDF")
    return buf.getvalue()


def test_docx_extracts_structure():
    res = LightParser().parse(_docx_bytes(), "docx", scanned_char_per_page_max=50)
    assert res.ok
    texts = [b.text for b in res.blocks if b.type is not BlockType.TABLE]
    assert any("第一章" in t for t in texts)
    assert any("第一条" in t for t in texts)
    assert all(b.page is None for b in res.blocks)  # docx page 待 B4 对齐回填
    tables = [b for b in res.blocks if b.type is BlockType.TABLE]
    assert len(tables) == 1
    assert tables[0].table.n_rows == 2 and tables[0].table.n_cols == 2
    assert res.title == "某单位综合管理办法"


def test_scanned_pdf_quarantined():
    res = LightParser().parse(_scanned_pdf_bytes(), "pdf", scanned_char_per_page_max=50)
    assert not res.ok and res.error_code == "E202-DEMO"  # 扫描件隔离


def test_unsupported_format():
    res = LightParser().parse(b"PK\x03\x04demo", "xlsx", scanned_char_per_page_max=50)
    assert not res.ok and res.error_code == "E101-DEMO"  # 白名单外


def test_pdf_text_layer(tmp_path):
    from pipeline.parsing.rendition import render_pdf, soffice_bin

    try:
        soffice_bin()
    except RuntimeError:
        pytest.skip("soffice 不可用")
    src = tmp_path / "a.docx"
    src.write_bytes(_docx_bytes())
    pdf = render_pdf(src, tmp_path)
    res = LightParser().parse(pdf.read_bytes(), "pdf", scanned_char_per_page_max=50)
    assert res.ok and res.blocks and res.page_count >= 1
    assert all(b.page is not None for b in res.blocks)  # pdf 原生页码
    assert any("第一条" in b.text for b in res.blocks)
