"""B 模式(``auto_confirm_meta_no_conflict``)端到端:无冲突新件 ingest 当场直达 INDEXED + finalize。

这是 B 模式"半接通" bug 的回归闸:旧实现下 ingest 用轻上下文 → 干净件流到 EMBEDDING 后**搁浅**
(无 s5 stage、又非 META_REVIEW,没有命令能捞它);即便接上 s5,run_until_idle 到 INDEXED 也**不调
finalize**(版本切换 + T2/T4 留痕全跳过)。本测试断言:无 ``_approve_doc`` 人工放行,干净新件经
``cli._drive_batch`` 自动到 INDEXED、跳过 meta_confirm 队列、且 finalize 留痕在位。

gate:PIPELINE_EMBEDDING_MODEL + PG + Milvus + soffice(同 T2 冒烟)。
"""

import os

import pytest
from docx import Document as Docx
from openpyxl import Workbook
from sqlalchemy import delete, select, text
from ulid import ULID

from common.pg_models import (
    Chunk,
    ClauseTag,
    Document,
    DocVersion,
    ImportBatch,
    PipelineEvent,
    RemediationRecord,
    ReviewQueue,
)
from pipeline import cli
from pipeline.config import load_config
from pipeline.index.embedding_client import EmbeddingClient
from pipeline.index.milvus_io import MilvusIO
from pipeline.index.object_store import ObjectStore
from pipeline.index.pg_io import PgIO
from pipeline.stage_base import StageContext
from pipeline.stages.s0_register import register_batch

_MANIFEST_COLS = [
    "filename", "title", "doc_number", "issuer", "perm_tag",
    "corpus_type", "biz_domain", "issue_date", "supersedes", "sub_type", "effective_date",
]


def _clean_docx(tmp_path):
    """唯一**无冲突**新件:首段 = manifest title(避开 ``unique_docx`` 的"首段≠manifest标题"冲突),
    body 无可抽文号/日期 → L1 交叉校验零冲突 → B 模式应自动放行。返回 ``(batch_dir, manifest)``。
    """
    tag = str(ULID())
    d = tmp_path / ("c_" + tag[:8])
    d.mkdir()
    fn, title = "clean.docx", "合同管理办法"
    doc = Docx()
    doc.add_paragraph(title)  # 首段即标题 → ir.title 与 manifest 一致(无 title 冲突)
    doc.add_paragraph("第一章 总则")
    doc.add_paragraph(
        f"第一条 为加强本单位合同管理规范合同签订与履行流程根据有关规定制定本办法编号{tag}。"
    )
    doc.add_paragraph("第二条 本办法适用于本单位各部门及全体人员的合同签订与履行活动。")
    doc.add_paragraph("第三条 合同应当经法务审查并由授权人签署后方可对外签订生效。")
    doc.save(d / fn)
    wb = Workbook()
    wb.active.append(_MANIFEST_COLS)
    # doc_number 设值但 body 无可抽文号 → meta.doc_numbers 空 → 不冲突(同 unique_docx 套路)
    dn = f"测试第{tag[:6]}号"
    wb.active.append(
        [fn, title, dn, "INTERNAL", "内部", "P-INT", "LEGAL", None, None, "内规", None]
    )
    mp = d / "manifest.xlsx"
    wb.save(mp)
    return d, mp


@pytest.fixture(scope="module")
def b_stack():
    """PG + Milvus + embedding 全栈,且**强制 B 模式**(toggle on,不依赖 settings.toml 取值)。"""
    if not os.environ.get("PIPELINE_EMBEDDING_MODEL"):
        pytest.skip("未设 PIPELINE_EMBEDDING_MODEL;B 模式端到端跳过")
    os.environ.setdefault("HF_HUB_OFFLINE", "1")
    os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")
    cfg = load_config()
    cfg = cfg.model_copy(
        update={"toggles": cfg.toggles.model_copy(update={"auto_confirm_meta_no_conflict": True})}
    )
    pg = PgIO.from_config(cfg)
    try:
        with pg.session() as s:
            s.execute(text("select 1"))
    except Exception:
        pytest.skip("PG 不可达")
    mio = MilvusIO(cfg)
    try:
        mio.connect()
        mio.create_collection()
    except Exception:
        pytest.skip("Milvus 不可达")
    emb = EmbeddingClient.from_config(cfg)
    try:
        emb.embed(["探测"])
    except Exception as e:
        pytest.skip(f"BGE-M3 加载失败: {e}")
    ctx = StageContext(
        config=cfg, object_store=ObjectStore.from_config(cfg), db=pg, embedding=emb, milvus=mio
    )
    yield pg, mio, ctx
    mio.disconnect()


def test_b_mode_clean_doc_auto_flows_to_indexed(b_stack, soffice, tmp_path):
    pg, mio, ctx = b_stack
    d, m = _clean_docx(tmp_path)  # 唯一无冲突新件(无 supersedes、首段=标题 → L1 零冲突)
    bid = str(ULID())
    register_batch(ctx, bid, d, m)
    dvids: list[str] = []
    try:
        cli._drive_batch(pg, ctx, bid)  # B 模式:无 _approve_doc,应自动到 INDEXED + finalize 扫尾
        with pg.session() as s:
            dvids = [
                x.doc_version_id
                for x in s.scalars(select(DocVersion).where(DocVersion.batch_id == bid))
            ]
        assert dvids
        (dvid,) = dvids
        # ① 无人工放行即到 INDEXED(旧 bug:会搁浅在 EMBEDDING)
        assert pg.get(DocVersion, dvid).pipeline_status == "INDEXED"
        with pg.session() as s:
            # ② 跳过 META_REVIEW → 无 open meta_confirm 队列项(B 模式自动放行)
            open_mc = list(
                s.scalars(
                    select(ReviewQueue)
                    .where(ReviewQueue.doc_version_id == dvid)
                    .where(ReviewQueue.queue_type == "meta_confirm")
                    .where(ReviewQueue.status == "open")
                )
            )
            assert open_mc == []
            evs = list(
                s.scalars(
                    select(PipelineEvent)
                    .where(PipelineEvent.doc_version_id == dvid)
                    .order_by(PipelineEvent.id)
                )
            )
        # ③ s4 自动放行留痕(STRUCTURING→EMBEDDING 带 auto_confirmed)
        assert any((e.detail or {}).get("auto_confirmed") for e in evs), "缺 auto_confirmed 留痕"
        # ④ finalize 扫尾跑了 T2/T4 并留痕(旧 bug:run_until_idle 不调 finalize → 无此痕)
        verify = [e.detail["verify"] for e in evs if (e.detail or {}).get("verify")]
        assert verify, "finalize 未留 verify 痕(扫尾未触发)"
        assert verify[-1]["t4_pass"] is True
    finally:
        _cleanup(pg, mio, bid, dvids)


def _cleanup(pg, mio, bid, dvids):
    for d_ in dvids:
        mio.delete(d_)
    mio.flush()
    with pg.session() as s:
        for d_ in dvids:
            s.execute(delete(RemediationRecord).where(RemediationRecord.doc_version_id == d_))
            s.execute(delete(ReviewQueue).where(ReviewQueue.doc_version_id == d_))
            s.execute(  # E1 clause_tags 是 chunk 的 FK 子,先删
                delete(ClauseTag).where(
                    ClauseTag.chunk_id.in_(
                        select(Chunk.chunk_id).where(Chunk.doc_version_id == d_)
                    )
                )
            )
            s.execute(delete(Chunk).where(Chunk.doc_version_id == d_))
            s.execute(delete(PipelineEvent).where(PipelineEvent.doc_version_id == d_))
        lids = [
            x.logical_id for x in s.scalars(select(DocVersion).where(DocVersion.batch_id == bid))
        ]
        s.execute(delete(DocVersion).where(DocVersion.batch_id == bid))
        if lids:
            s.execute(delete(Document).where(Document.logical_id.in_(lids)))
        s.execute(delete(ImportBatch).where(ImportBatch.batch_id == bid))
