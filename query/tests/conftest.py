"""查询集成测试共享栈:连真 PG+Milvus+BGE-M3,ingest 一件内规到 INDEXED(B 模式自动放行)。

gate:PIPELINE_EMBEDDING_MODEL + PG + Milvus + soffice(同摄取侧 B 模式端到端)。任一不满足即 skip。
session 作用域:整批查询集成测试共用同一件 INDEXED 内规,结束反 FK 序清理 + 清 Milvus 投影。
"""

from __future__ import annotations

import os
from collections import namedtuple

import pytest
from docx import Document as Docx
from openpyxl import Workbook
from sqlalchemy import delete, select, text
from ulid import ULID

from common.pg_models import (
    Case,
    Chunk,
    ClauseTag,
    Document,
    DocVersion,
    ImportBatch,
    PipelineEvent,
    RemediationRecord,
    ReviewQueue,
)
from pipeline import cli
from pipeline.config import load_config
from pipeline.index.embedding_client import EmbeddingClient
from pipeline.index.milvus_io import MilvusIO
from pipeline.index.object_store import ObjectStore
from pipeline.index.pg_io import PgIO
from pipeline.stage_base import StageContext
from pipeline.stages.s0_register import register_batch

_MANIFEST_COLS = [
    "filename",
    "title",
    "doc_number",
    "issuer",
    "perm_tag",
    "corpus_type",
    "biz_domain",
    "issue_date",
    "supersedes",
    "sub_type",
    "effective_date",
]

#: 稳定查询词(命中 ingest 件第三条「合同应当经法务审查并由授权人签署」)
QUERY_TEXT = "合同应当经法务审查并由授权人签署"
#: R3 案例问句(classify → CASE scene:含「处罚案例」关键词)
CASE_QUERY = "有没有微信二维码违规开户的处罚案例"

IndexedStack = namedtuple("IndexedStack", "pg mio ctx dvid query")
#: R3 案例栈(复用 indexed_stack 的内规件 + 额外一件 P-CASE 处罚决定书)
CaseStack = namedtuple("CaseStack", "pg mio ctx internal_dvid case_dvid query case_query")
#: R4 列举栈(复用 indexed_stack + 额外两件同主题「信息披露」内规:a 含义务条款、b 无义务)
EnumStack = namedtuple("EnumStack", "pg mio ctx dvid_a dvid_b biz_code")
#: §5.4 sparse 提权/扩展栈(复用 indexed_stack + 一件:发文字号条款 + 合同竞争条款 + 受托理财条款)
SparseStack = namedtuple("SparseStack", "pg mio ctx dvid docnum_query oral_query")


def _clean_internal_docx(tmp_path):
    """唯一无冲突内规件(首段=manifest 标题、body 无可抽文号 → L1 零冲突 → B 模式自动放行)。"""
    tag = str(ULID())
    d = tmp_path / ("q_" + tag[:8])
    d.mkdir()
    fn, title = "clean.docx", "合同管理办法"
    doc = Docx()
    doc.add_paragraph(title)
    doc.add_paragraph("第一章 总则")
    doc.add_paragraph("第一节 一般规定")  # 章→节→条:产出节级父块,使 §5.6 父块供证可验
    doc.add_paragraph(
        f"第一条 为加强本单位合同管理规范合同签订与履行流程根据有关规定制定本办法编号{tag}。"
    )
    doc.add_paragraph("第二条 本办法适用于本单位各部门及全体人员的合同签订与履行活动。")
    doc.add_paragraph("第二节 签订与履行")
    doc.add_paragraph("第三条 合同应当经法务审查并由授权人签署后方可对外签订生效并妥善归档备查。")
    doc.save(d / fn)
    wb = Workbook()
    wb.active.append(_MANIFEST_COLS)
    wb.active.append(
        [
            fn,
            title,
            f"测试第{tag[:6]}号",
            "INTERNAL",
            "内部",
            "P-INT",
            "LEGAL",
            None,
            None,
            "内规",
            None,
        ]
    )
    mp = d / "manifest.xlsx"
    wb.save(mp)
    return d, mp


def _penalty_case_docx(tmp_path):
    """唯一无冲突 P-CASE 处罚决定书:首段=manifest 标题、body 无可抽文号 → L1 零冲突 → B 模式放行。

    L1 案例要素(case_extract):penalty_org=北京证监局(头部抬头)/ respondent / penalty_date /
    penalty_type / amount_wan;cited_regulations 为 L2 字段默认空(精确反查由集成测手插验证)。
    """
    tag = str(ULID())
    d = tmp_path / ("qc_" + tag[:8])
    d.mkdir()
    fn, title = "penalty.docx", "北京证监局行政处罚决定书"
    doc = Docx()
    doc.add_paragraph(title)  # 首段=manifest 标题(避免 title 冲突)
    doc.add_paragraph("当事人:某某证券有限公司,住所地北京市朝阳区。")
    doc.add_paragraph(
        f"经查,该公司存在以下违规行为:通过微信发送开户推广二维码违规招揽客户编号{tag}。"
    )
    doc.add_paragraph("处罚依据:依据《证券法》第一百九十七条的规定。")
    doc.add_paragraph("现决定:对当事人给予警告,并处以罚款50万元。")
    doc.add_paragraph("2024年3月15日")
    doc.save(d / fn)
    wb = Workbook()
    wb.active.append(_MANIFEST_COLS)
    wb.active.append(
        [
            fn,
            title,
            f"案例第{tag[:6]}号",  # manifest 文号;body 无可抽文号 → meta.doc_numbers 空 → 不冲突
            "INTERNAL",            # issuer:不解析到 dict code → 无 issuer 冲突(同内规件)
            "内部",
            "P-CASE",
            "LEGAL",
            None,                  # issue_date None → 不与 body 日期冲突
            None,
            "案例",
            None,
        ]
    )
    mp = d / "manifest.xlsx"
    wb.save(mp)
    return d, mp


@pytest.fixture(scope="session")
def indexed_stack(soffice, tmp_path_factory):
    if not os.environ.get("PIPELINE_EMBEDDING_MODEL"):
        pytest.skip("未设 PIPELINE_EMBEDDING_MODEL;查询集成跳过")
    os.environ.setdefault("HF_HUB_OFFLINE", "1")
    os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")
    cfg = load_config()
    cfg = cfg.model_copy(
        update={"toggles": cfg.toggles.model_copy(update={"auto_confirm_meta_no_conflict": True})}
    )
    pg = PgIO.from_config(cfg)
    try:
        with pg.session() as s:
            s.execute(text("select 1"))
    except Exception:
        pytest.skip("PG 不可达(demo up 未起)")
    mio = MilvusIO(cfg)
    try:
        mio.connect()
        mio.create_collection()
    except Exception:
        pytest.skip("Milvus 不可达")
    emb = EmbeddingClient.from_config(cfg)
    try:
        emb.embed(["探测"])
    except Exception as e:
        pytest.skip(f"BGE-M3 加载失败: {e}")
    ctx = StageContext(
        config=cfg, object_store=ObjectStore.from_config(cfg), db=pg, embedding=emb, milvus=mio
    )

    tmp_path = tmp_path_factory.mktemp("q_ingest")
    d, m = _clean_internal_docx(tmp_path)
    bid = str(ULID())
    register_batch(ctx, bid, d, m)
    cli._drive_batch(pg, ctx, bid)  # B 模式:无人工放行自动到 INDEXED + finalize
    with pg.session() as s:
        dvids = [
            x.doc_version_id
            for x in s.scalars(select(DocVersion).where(DocVersion.batch_id == bid))
        ]
    assert dvids, "ingest 未产出 dvid"
    (dvid,) = dvids
    dv = pg.get(DocVersion, dvid)
    assert dv.pipeline_status == "INDEXED", f"未到 INDEXED:{dv.pipeline_status}"
    logical_id = dv.logical_id

    yield IndexedStack(pg, mio, ctx, dvid, QUERY_TEXT)

    # 反 FK 序清理 + Milvus 投影
    try:
        mio.delete(dvid)
        mio.flush()
    except Exception:
        pass
    with pg.session() as s:
        child_ids = select(Chunk.chunk_id).where(Chunk.doc_version_id == dvid)
        s.execute(delete(ClauseTag).where(ClauseTag.chunk_id.in_(child_ids)))
        s.execute(delete(Chunk).where(Chunk.doc_version_id == dvid))
        s.execute(delete(PipelineEvent).where(PipelineEvent.doc_version_id == dvid))
        s.execute(delete(RemediationRecord).where(RemediationRecord.doc_version_id == dvid))
        s.execute(delete(ReviewQueue).where(ReviewQueue.doc_version_id == dvid))
        s.execute(delete(DocVersion).where(DocVersion.doc_version_id == dvid))
        s.execute(delete(Document).where(Document.logical_id == logical_id))
        s.execute(delete(ImportBatch).where(ImportBatch.batch_id == bid))
    mio.disconnect()


#: R4 列举栈共享业务域(biz_domain code,Milvus ARRAY 存此值;供 extra_expr biz 过滤验证)
ENUM_BIZ_CODE = "DISCLOSURE"


def _disclosure_docx(tmp_path, prefix, title, clauses):
    """同主题「信息披露」内规件(首段=manifest 标题、body 无可抽文号 → L1 零冲突 → B 模式放行)。

    ``clauses`` 为「第N条…」正文行列表;``ENUM_BIZ_CODE`` 入 manifest biz_domain(→ Milvus ARRAY)。
    """
    tag = str(ULID())
    d = tmp_path / (prefix + tag[:8])
    d.mkdir()
    fn = "disc.docx"
    doc = Docx()
    doc.add_paragraph(title)            # 首段=manifest 标题(免 title 冲突)
    doc.add_paragraph("第一章 总则")
    doc.add_paragraph("第一节 一般规定")  # 章→节→条:产出节级父块
    for i, body in enumerate(clauses):
        doc.add_paragraph(f"{body}编号{tag}{i}。")
    doc.save(d / fn)
    wb = Workbook()
    wb.active.append(_MANIFEST_COLS)
    wb.active.append(
        [fn, title, f"披露第{tag[:6]}号", "INTERNAL", "内部", "P-INT",
         ENUM_BIZ_CODE, None, None, "内规", None]
    )
    mp = d / "manifest.xlsx"
    wb.save(mp)
    return d, mp


def _ingest_one(pg, ctx, d, m):
    """ingest 一件到 INDEXED(B 模式自动放行),返回 (dvid, logical_id, batch_id)。"""
    bid = str(ULID())
    register_batch(ctx, bid, d, m)
    cli._drive_batch(pg, ctx, bid)
    with pg.session() as s:
        (dvid,) = [
            x.doc_version_id
            for x in s.scalars(select(DocVersion).where(DocVersion.batch_id == bid))
        ]
    dv = pg.get(DocVersion, dvid)
    assert dv.pipeline_status == "INDEXED", f"未到 INDEXED:{dv.pipeline_status}"
    return dvid, dv.logical_id, bid


def _purge_doc(pg, mio, dvid, logical_id, bid):
    """反 FK 序清理一件 + Milvus 投影。"""
    try:
        mio.delete(dvid)
        mio.flush()
    except Exception:
        pass
    with pg.session() as s:
        child_ids = select(Chunk.chunk_id).where(Chunk.doc_version_id == dvid)
        s.execute(delete(ClauseTag).where(ClauseTag.chunk_id.in_(child_ids)))
        s.execute(delete(Chunk).where(Chunk.doc_version_id == dvid))
        s.execute(delete(PipelineEvent).where(PipelineEvent.doc_version_id == dvid))
        s.execute(delete(RemediationRecord).where(RemediationRecord.doc_version_id == dvid))
        s.execute(delete(ReviewQueue).where(ReviewQueue.doc_version_id == dvid))
        s.execute(delete(DocVersion).where(DocVersion.doc_version_id == dvid))
        s.execute(delete(Document).where(Document.logical_id == logical_id))
        s.execute(delete(ImportBatch).where(ImportBatch.batch_id == bid))


@pytest.fixture(scope="session")
def enumerate_stack(indexed_stack, tmp_path_factory):
    """R4:复用 indexed_stack + 额外 ingest 两件同主题「信息披露」内规到 INDEXED。

    doc_a 第二条含义务条款(``应当`` → E1 自动打 ``is_obligation``);doc_b 全无义务标记
    → 义务查询时 doc_b 被 E1 后过滤剔除,doc_a 保留(验 consumed-when-present 过滤)。
    """
    pg, mio, ctx = indexed_stack.pg, indexed_stack.mio, indexed_stack.ctx
    tmp = tmp_path_factory.mktemp("q_enum")
    da, ma = _disclosure_docx(
        tmp, "enum_a_", "信息披露管理办法",
        ["第一条 为规范信息披露根据有关规定制定本办法",
         "第二条 上市公司应当及时完整披露信息披露相关重大事项"],  # 第二条:应当→义务
    )
    db, mb = _disclosure_docx(
        tmp, "enum_b_", "信息披露事务管理细则",
        ["第一条 本细则界定信息披露的范围与办理流程",
         "第二条 信息披露分为定期报告与临时公告两类内容"],  # 无义务标记
    )
    dvid_a, lid_a, bid_a = _ingest_one(pg, ctx, da, ma)
    dvid_b, lid_b, bid_b = _ingest_one(pg, ctx, db, mb)

    yield EnumStack(pg, mio, ctx, dvid_a, dvid_b, ENUM_BIZ_CODE)

    _purge_doc(pg, mio, dvid_b, lid_b, bid_b)
    _purge_doc(pg, mio, dvid_a, lid_a, bid_a)


#: §5.4 发文字号查询(含发文字号 + 语义「合同管理」→ 不提权时 dense 易偏向合同竞争条款)
SPARSE_DOCNUM_QUERY = "银保监发〔2021〕5号 合同管理要求"
#: §5.4 口语查询(dict 映射 代客理财→受托理财;条款只含法言词「受托理财」)
SPARSE_ORAL_QUERY = "代客理财是否违规"


#: §5.4 发文字号(嵌第一条正文供提权命中);manifest doc_number 设同值 → L1 无冲突
SPARSE_DOCNUM = "银保监发〔2021〕5号"


def _sparse_docx(tmp_path):
    """§5.4 件:第一条含发文字号(冒号边界 → L1 抽取与 manifest 一致);第二条合同管理(竞争块);
    第三条含法言词「受托理财」(无「代客理财」)。首段=manifest 标题 → B 模式放行。
    """
    tag = str(ULID())
    d = tmp_path / ("qs_" + tag[:8])
    d.mkdir()
    fn, title = "sparse.docx", "合规管理办法"
    doc = Docx()
    doc.add_paragraph(title)
    doc.add_paragraph("第一章 总则")
    doc.add_paragraph("第一节 一般规定")
    # 冒号边界 → 文号正则前缀只吃「银保监发」→ 抽取=SPARSE_DOCNUM=manifest doc_number(无冲突)
    doc.add_paragraph(f"第一条 本条适用文号:{SPARSE_DOCNUM},具体编号{tag}。")
    doc.add_paragraph(f"第二条 合同管理要求合同应当经法务审查后由授权人签署编号{tag}2。")
    doc.add_paragraph(f"第三条 公司不得以受托理财名义违规开展资产管理业务编号{tag}3。")
    doc.save(d / fn)
    wb = Workbook()
    wb.active.append(_MANIFEST_COLS)
    wb.active.append(
        [fn, title, SPARSE_DOCNUM, "INTERNAL", "内部", "P-INT",
         "LEGAL", None, None, "内规", None]
    )
    mp = d / "manifest.xlsx"
    wb.save(mp)
    return d, mp


@pytest.fixture(scope="session")
def sparse_stack(indexed_stack, tmp_path_factory):
    """§5.4:复用 indexed_stack 真栈 + 额外 ingest 一件(发文字号/合同/受托理财 三条)到 INDEXED。"""
    pg, mio, ctx = indexed_stack.pg, indexed_stack.mio, indexed_stack.ctx
    tmp = tmp_path_factory.mktemp("q_sparse")
    d, m = _sparse_docx(tmp)
    dvid, lid, bid = _ingest_one(pg, ctx, d, m)
    yield SparseStack(pg, mio, ctx, dvid, SPARSE_DOCNUM_QUERY, SPARSE_ORAL_QUERY)
    _purge_doc(pg, mio, dvid, lid, bid)


@pytest.fixture(scope="session")
def case_stack(indexed_stack, tmp_path_factory):
    """R3:复用 indexed_stack 内规件 + 额外 ingest 一件 P-CASE 处罚决定书到 INDEXED。"""
    pg, mio, ctx = indexed_stack.pg, indexed_stack.mio, indexed_stack.ctx
    tmp = tmp_path_factory.mktemp("q_case")
    d, m = _penalty_case_docx(tmp)
    bid = str(ULID())
    register_batch(ctx, bid, d, m)
    cli._drive_batch(pg, ctx, bid)  # B 模式自动到 INDEXED
    with pg.session() as s:
        cdvids = [
            x.doc_version_id
            for x in s.scalars(select(DocVersion).where(DocVersion.batch_id == bid))
        ]
    assert cdvids, "案例件 ingest 未产出 dvid"
    (case_dvid,) = cdvids
    dv = pg.get(DocVersion, case_dvid)
    assert dv.pipeline_status == "INDEXED", f"案例件未到 INDEXED:{dv.pipeline_status}"
    assert pg.get_case(case_dvid) is not None, "cases 表未回填案例要素"
    clogical = dv.logical_id

    yield CaseStack(pg, mio, ctx, indexed_stack.dvid, case_dvid, indexed_stack.query, CASE_QUERY)

    # 反 FK 序清理案例件(cases 行先于 doc_versions)+ Milvus 投影
    try:
        mio.delete(case_dvid)
        mio.flush()
    except Exception:
        pass
    with pg.session() as s:
        child_ids = select(Chunk.chunk_id).where(Chunk.doc_version_id == case_dvid)
        s.execute(delete(ClauseTag).where(ClauseTag.chunk_id.in_(child_ids)))
        s.execute(delete(Case).where(Case.doc_version_id == case_dvid))
        s.execute(delete(Chunk).where(Chunk.doc_version_id == case_dvid))
        s.execute(delete(PipelineEvent).where(PipelineEvent.doc_version_id == case_dvid))
        s.execute(delete(RemediationRecord).where(RemediationRecord.doc_version_id == case_dvid))
        s.execute(delete(ReviewQueue).where(ReviewQueue.doc_version_id == case_dvid))
        s.execute(delete(DocVersion).where(DocVersion.doc_version_id == case_dvid))
        s.execute(delete(Document).where(Document.logical_id == clogical))
        s.execute(delete(ImportBatch).where(ImportBatch.batch_id == bid))
